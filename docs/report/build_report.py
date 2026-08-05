"""Build the technical report: Markdown -> DOCX -> PDF.

    python docs/report/build_report.py

Stage 1 (pandoc) renders `technical_report.md` into `Group25_TechnicalReport.docx`, using a
reference document generated here so the output meets the assignment's formatting rules
(Times New Roman, 11 pt body, 0.9-inch margins, numbered figure/table captions come from the
Markdown itself).

Stage 2 exports that DOCX to PDF through Word (COM automation, Windows only). If Word is not
available the DOCX is still produced and the script says so instead of failing the build.

Requires: pandoc on PATH, python-docx. Word is optional (PDF stage only).
"""
from __future__ import annotations

import os
import shutil
import subprocess
import sys

try:                                          # only needed for the reference document
    from docx.oxml.ns import qn
except ImportError:                           # python-docx missing; main() reports it
    qn = None

HERE = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.dirname(os.path.dirname(HERE))
SOURCE = os.path.join(HERE, "technical_report.md")
REFERENCE = os.path.join(HERE, "reference.docx")
DOCX = os.path.join(HERE, "Group25_TechnicalReport.docx")
PDF = os.path.join(HERE, "Group25_TechnicalReport.pdf")

BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 11


def find_style(document, name: str):
    """Look a style up by its literal name.

    `document.styles[name]` translates built-in UI names to python-docx's internal spelling
    ("Heading 1" -> "heading 1"), which does not match how pandoc's reference document spells
    them, so every heading lookup would raise KeyError. Matching the raw name avoids that.
    """
    return next((s for s in document.styles if s.name == name), None)


def force_font(style, name: str) -> None:
    """Pin a style's font, overriding the theme fonts pandoc's headings inherit.

    Setting `style.font.name` only writes `w:rFonts/@ascii`; Word still prefers the
    `asciiTheme`/`hAnsiTheme` attributes if they are present, which is why headings render
    sans-serif otherwise.
    """
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        if rfonts.get(qn(f"w:{attr}")) is not None:
            del rfonts.attrib[qn(f"w:{attr}")]
    for attr in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def build_reference_docx(path: str = REFERENCE) -> str:
    """Write the pandoc reference document that carries the required formatting."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    default = os.path.join(HERE, "_pandoc_default.docx")
    subprocess.run(["pandoc", "--print-default-data-file", "reference.docx"],
                   stdout=open(default, "wb"), check=True)

    doc = Document(default)
    for section in doc.sections:
        # Uniform 0.9 in margins: keeps the report inside the 8-10 page budget without
        # dropping body text below the required 11 pt.
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(0.9)

    normal = find_style(doc, "Normal")
    force_font(normal, BODY_FONT)
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, bold in (("Title", 16, True), ("Heading 1", 13, True),
                             ("Heading 2", 11.5, True), ("Heading 3", 11, True),
                             ("Author", 11, False), ("Date", 10, False),
                             ("Body Text", BODY_SIZE_PT, False),
                             ("First Paragraph", BODY_SIZE_PT, False),
                             ("Compact", BODY_SIZE_PT, False)):
        style = find_style(doc, name)
        if style is not None:
            force_font(style, BODY_FONT)
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = RGBColor(0, 0, 0)

    # Headings default to generous space; the page budget is better spent on content.
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = find_style(doc, name)
        if style is not None:
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(3)

    for name in ("Caption", "Image Caption", "Table Caption"):
        style = find_style(doc, name)
        if style is not None:
            style.font.size = Pt(9)
            style.font.italic = True
            style.font.name = BODY_FONT
            # Captions sit *after* their figure/table here, so keep-with-next would bind them
            # to the following paragraph and push whole blocks to the next page — which was
            # costing ~3 inches of whitespace per occurrence.
            style.paragraph_format.keep_with_next = False

    # pandoc's Body Text carries a 9 pt space-after that overrides Normal's; over ~45
    # paragraphs that alone is most of a page.
    for name in ("Body Text", "First Paragraph", "Compact", "Block Text"):
        style = find_style(doc, name)
        if style is not None:
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.0

    # Widow/orphan control reserves a spare line on every page and, combined with the
    # keep-with-next chains above, was leaving ~3 inches unused on figure pages.
    for name in ("Normal", "Body Text", "First Paragraph", "Compact", "Block Text",
                 "Heading 1", "Heading 2", "Heading 3"):
        style = find_style(doc, name)
        if style is not None:
            style.paragraph_format.widow_control = False

    doc.save(path)
    os.remove(default)
    return path


def markdown_to_docx() -> str:
    """Render the Markdown source to DOCX with pandoc."""
    subprocess.run(
        # No --toc: the assignment's 8-10 page budget is better spent on results.
        ["pandoc", SOURCE, "--from", "markdown+pipe_tables+implicit_figures",
         "--reference-doc", REFERENCE, "--resource-path", REPO_ROOT,
         "--output", DOCX],
        cwd=REPO_ROOT, check=True)
    return DOCX


def tighten_tables(path: str = DOCX, size_pt: float = 9.5) -> str:
    """Shrink table text only.

    pandoc styles table cells and tight list items with the same `Compact` style, so the
    table size cannot be set in the reference document without also dropping bulleted list
    items below the assignment's 11 pt body-text floor. Setting the runs directly after the
    fact keeps prose at 11 pt and still fits nine result tables in the page budget.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(size_pt)
    doc.save(path)
    return path


def docx_to_pdf() -> str | None:
    """Export the DOCX to PDF through Word; returns None when Word is unavailable."""
    try:
        import win32com.client                                   # pywin32
    except ImportError:
        print("[pdf] pywin32 not installed — skipping PDF export")
        return None

    word = None
    try:
        # EnsureDispatch (early binding) — late binding drops the connection mid-export here.
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(DOCX, ConfirmConversions=False, ReadOnly=True,
                                  AddToRecentFiles=False)
        doc.ExportAsFixedFormat(OutputFileName=PDF, ExportFormat=17,      # wdExportFormatPDF
                                OpenAfterExport=False)
        pages = doc.ComputeStatistics(2)                                  # wdStatisticPages
        doc.Close(0)
        print(f"[pdf] {os.path.basename(PDF)} — {pages} pages")
        return PDF
    except Exception as exc:                                     # Word missing or COM refused
        print(f"[pdf] Word export failed ({exc}) — the DOCX is still valid")
        return None
    finally:
        if word is not None:
            word.Quit()


def main() -> int:
    if not shutil.which("pandoc"):
        print("pandoc is required: https://pandoc.org/installing.html")
        return 1
    if not os.path.exists(SOURCE):
        print(f"missing source: {SOURCE}")
        return 1

    build_reference_docx()
    markdown_to_docx()
    tighten_tables()
    print(f"[docx] {os.path.basename(DOCX)}")
    docx_to_pdf()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
